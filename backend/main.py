"""
화성시 계약업무 시스템 — 백엔드
실행: python backend/main.py
"""

import io
import re
import zipfile
import json

from pathlib import Path
from fastapi import FastAPI, File, UploadFile, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse

app = FastAPI(title="화성시 계약업무 API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── 경로 설정 ─────────────────────────────────────────────
BASE_DIR = Path(__file__).parent.parent
DATA_FILE = BASE_DIR / "backend" / "data" / "contract_data.json"
DATA_FILE.parent.mkdir(exist_ok=True)


# ── 데이터 읽기 / 쓰기 ────────────────────────────────────
def read_data() -> dict:
    if DATA_FILE.exists():
        return json.loads(DATA_FILE.read_text(encoding="utf-8"))
    return {}


def write_data(data: dict):
    DATA_FILE.write_text(
        json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8"
    )


# ── 추출 정규식 패턴 ──────────────────────────────────────
PATTERNS = {
    "amount": [
        r"금\s*([\d,]+)\s*원",
        r"계약금액\s*[:：]\s*([\d,]+)",
        r"금\s*액\s*[:：]\s*([\d,]+)",
        r"총\s*금\s*액\s*[:：]?\s*([\d,]+)",
    ],
    "biz_num": [
        r"\d{3}-\d{2}-\d{5}",
    ],
    "start_date": [
        r"계약\s*일\s*[:：]?\s*(\d{4}[년.\-]\s*\d{1,2}[월.\-]\s*\d{1,2})",
        r"착공\s*일\s*[:：]?\s*(\d{4}[년.\-]\s*\d{1,2}[월.\-]\s*\d{1,2})",
        r"시작\s*일\s*[:：]?\s*(\d{4}[년.\-]\s*\d{1,2}[월.\-]\s*\d{1,2})",
    ],
    "end_date": [
        r"준공\s*일\s*[:：]?\s*(\d{4}[년.\-]\s*\d{1,2}[월.\-]\s*\d{1,2})",
        r"납품\s*기한\s*[:：]?\s*(\d{4}[년.\-]\s*\d{1,2}[월.\-]\s*\d{1,2})",
        r"완료\s*일\s*[:：]?\s*(\d{4}[년.\-]\s*\d{1,2}[월.\-]\s*\d{1,2})",
        r"납품\s*일\s*[:：]?\s*(\d{4}[년.\-]\s*\d{1,2}[월.\-]\s*\d{1,2})",
    ],
    "vendor": [
        r"(?:상\s*호|업체\s*명|회사\s*명|수급\s*인|공급\s*자)\s*[:：]\s*(.+?)(?=\s*(?:대표|사업자|주소|$))",
        r"(?:상\s*호|업체\s*명|회사\s*명)\s*[:：]\s*(.{2,50})",
    ],
    "rep": [
        r"대표\s*자\s*[:：]\s*(.+?)(?=\s*(?:사업자|주소|생년|$))",
        r"대표\s*자\s*[:：]\s*(.{1,20})",
    ],
    "address": [
        r"(?:주\s*소|사업장\s*주소|소\s*재\s*지)\s*[:：]\s*(.+?)(?=\s*(?:전화|TEL|연락|$))",
        r"(?:주\s*소|사업장\s*주소)\s*[:：]\s*(.{5,100})",
    ],
    "phone": [
        r"(?:전화|연락처|TEL|전화번호)\s*[:：]?\s*((?:0\d{1,2}[-\s]?\d{3,4}[-\s]?\d{4}))",
    ],
    "name": [
        r"(?:공사\s*명|용역\s*명|물품\s*명|사업\s*명|건\s*명)\s*[:：]\s*(.+?)(?=\s*(?:계약|금액|$))",
        r"(?:공사\s*명|용역\s*명|물품\s*명)\s*[:：]\s*(.{2,100})",
    ],
}


# ── 텍스트 추출 함수 ─────────────────────────────────────
def extract_text_pdf(content: bytes) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
    return "\n".join(parts)


def extract_text_hwpx(content: bytes) -> str:
    """HWPX: ZIP 내부의 XML 파일에서 텍스트 노드 추출"""
    parts = []
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as z:
            for name in z.namelist():
                lower = name.lower()
                if lower.endswith(".xml") and any(k in lower for k in ("section", "content", "body")):
                    try:
                        xml_bytes = z.read(name)
                        from lxml import etree
                        root = etree.fromstring(xml_bytes)
                        texts = [t.strip() for t in root.itertext() if t.strip()]
                        parts.append(" ".join(texts))
                    except Exception:
                        try:
                            raw = z.read(name).decode("utf-8", errors="ignore")
                            raw = re.sub(r"<[^>]+>", " ", raw)
                            parts.append(raw)
                        except Exception:
                            pass
    except zipfile.BadZipFile:
        pass
    return "\n".join(parts)


def extract_text_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text.strip())
    return "\n".join(lines)


# ── 필드 파싱 ────────────────────────────────────────────
def parse_fields(text: str) -> dict:
    result = {
        "name": "", "amount": "", "start_date": "", "end_date": "",
        "vendor": "", "rep": "", "biz_num": "", "address": "",
        "phone": "", "industry": "",
    }

    m = re.search(r"\d{3}-\d{2}-\d{5}", text)
    if m:
        result["biz_num"] = m.group(0)

    for pat in PATTERNS["amount"]:
        m = re.search(pat, text)
        if m:
            result["amount"] = m.group(1).replace(",", "")
            break

    for key in ("start_date", "end_date"):
        for pat in PATTERNS[key]:
            m = re.search(pat, text)
            if m:
                result[key] = m.group(1).strip()
                break

    for key in ("vendor", "rep", "address", "phone", "name"):
        for pat in PATTERNS[key]:
            m = re.search(pat, text, re.MULTILINE | re.DOTALL)
            if m:
                val = m.group(1).strip().split("\n")[0].strip()
                if len(val) >= 2:
                    result[key] = val[:120]
                    break

    if result["phone"]:
        result["phone"] = re.sub(r"[-\s]+", "-", result["phone"])

    result["extracted_count"] = sum(
        1 for k, v in result.items() if k != "extracted_count" and v
    )
    return result


# ── API: 계약 데이터 (JSON 파일 영속) ────────────────────
@app.get("/api/data")
def get_data():
    return JSONResponse(content=read_data())


@app.post("/api/data")
async def save_data(request: Request):
    body = await request.json()
    write_data(body)
    return {"status": "saved"}


@app.delete("/api/data")
def clear_data():
    write_data({})
    return {"status": "cleared"}


# ── API: 파일 자동 추출 ───────────────────────────────────
@app.post("/api/extract")
async def extract(file: UploadFile = File(...)):
    content = await file.read()
    filename = (file.filename or "").lower()

    try:
        if filename.endswith(".pdf"):
            text = extract_text_pdf(content)
        elif filename.endswith(".hwpx"):
            text = extract_text_hwpx(content)
        elif filename.endswith(".docx"):
            text = extract_text_docx(content)
        else:
            return {
                "name": "", "amount": "", "start_date": "", "end_date": "",
                "vendor": "", "rep": "", "biz_num": "", "address": "",
                "phone": "", "industry": "", "extracted_count": 0,
                "error": "지원하지 않는 파일 형식"
            }
        return parse_fields(text)

    except Exception as e:
        return {
            "name": "", "amount": "", "start_date": "", "end_date": "",
            "vendor": "", "rep": "", "biz_num": "", "address": "",
            "phone": "", "industry": "", "extracted_count": 0,
            "error": str(e)
        }


# ── 정적 파일 서빙 ───────────────────────────────────────
app.mount("/assets", StaticFiles(directory=str(BASE_DIR / "assets")), name="assets")
app.mount("/forms",  StaticFiles(directory=str(BASE_DIR / "forms")),  name="forms")
app.mount("/admin",  StaticFiles(directory=str(BASE_DIR / "admin"), html=True), name="admin")


@app.get("/", response_class=FileResponse)
def index():
    return FileResponse(str(BASE_DIR / "index.html"))


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=False)
